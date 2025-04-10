import json
import logging

from docx import Document

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.resume_model import Resume



class Parser:
    """ 
        Python class using pydantic to validate and parse a resume from json into 
        either a Markdown, Pdf, or WordDoc
    """

    def __init__(self, json_path: str = None, data:Resume = None) -> None:
        
        self.json_path = json_path
        if not json_path:
            raise ValueError("JSON path cannot be None")
        
        if not json_path.endswith('resume.json'):
            raise ValueError("File must be a JSON file")
        

        
        self.data = data
        if isinstance(self.data, Resume) and not None:
            self.resume = self.data
        
        if not self.data:
            self.__load_json()

    def __load_json(self) -> None:
        with open(self.json_path, 'r') as f:
            self.resume = Resume(**json.load(f))
            self.data = self.resume

    def __get_resume(self) -> Resume:
        return self.resume
    
    def get_markdown(self):

        with open("./docs/index.md", 'w') as f:

            f.write(f"# {self.resume.name}\n\n")
            f.write(f"**Phone:** {self.resume.contact.phone}  \n")
            f.write(f"**Email:** {self.resume.contact.email}  \n")
            f.write(f"**GitHub:** [{self.resume.contact.github}]({self.resume.contact.github})  \n")
            f.write(f"**LinkedIn:** [{self.resume.contact.linkedin}]({self.resume.contact.linkedin})\n\n")



            f.write(f"```\n")

            f.write(f" name: {self.resume.name}\n")
            f.write(f" phone: {self.resume.contact.phone}\n")
            f.write(f" email: {self.resume.contact.email}\n")
            f.write(f" github: {self.resume.contact.github}\n")
            f.write(f" linkedin: {self.resume.contact.linkedin}\n")
            f.write("```\n\n")

            f.write(f"## Summary\n\n")
            f.write(f"{self.resume.summary}\n\n")

            f.write("## Objective\n\n")
            f.write(f"{self.resume.objective}\n\n")

            
            
            f.write("## Experience\n\n")
            for job in self.resume.experience:
                f.write(f"### {job.company} ({job.location})\n")
                f.write(f"**{job.title}**\n")
                f.write(f"*{job.dates}*\n\n")
                for responsibility in job.responsibilities:
                    f.write(f"- {responsibility}\n")
                f.write("\n")
            
            f.write("## Skills\n\n")
            for skill_category, skills in self.resume.skills.model_dump().items():
                if skills:
                    category = skill_category.replace("_", " ")
                    f.write(f"### {category.capitalize()}\n")
                    for skill in skills:
                        f.write(f"- {skill}\n")
                    f.write("\n")

    def get_word_doc(self, output_path="./docs/resume.docx"):
        document = Document()

        # Add name and contact information
        document.add_heading(self.resume.name, level=1)
        document.add_paragraph(f"Phone: {self.resume.contact.phone}")
        document.add_paragraph(f"Email: {self.resume.contact.email}")
        document.add_paragraph(f"GitHub: {self.resume.contact.github}")
        document.add_paragraph(f"LinkedIn: {self.resume.contact.linkedin}")

        # Add summary
        document.add_heading("Summary", level=2)
        document.add_paragraph(self.resume.summary)

        # Add objective
        document.add_heading("Objective", level=2)
        document.add_paragraph(self.resume.objective)

        # Add experience
        document.add_heading("Experience", level=2)
        for job in self.resume.experience:
            document.add_heading(f"{job.company} ({job.location})", level=3)
            document.add_paragraph(f"{job.title} ({job.dates})")
            for responsibility in job.responsibilities:
                document.add_paragraph(f"- {responsibility}", style="List Bullet")

        # Add skills
        document.add_heading("Skills", level=2)
        for skill_category, skills in self.resume.skills.model_dump().items():
            if skills:
                document.add_heading(skill_category.replace("_", " ").capitalize(), level=3)
                for skill in skills:
                    document.add_paragraph(f"- {skill}", style="List Bullet")

        # Save the document
        document.save(output_path)

    def get_pdf(self, output_path="./docs/resume.pdf"):
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter

        # Add name and contact information
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, self.resume.name)

        c.setFont("Helvetica", 12)
        c.drawString(50, height - 80, f"Phone: {self.resume.contact.phone}")
        c.drawString(50, height - 100, f"Email: {self.resume.contact.email}")
        c.drawString(50, height - 120, f"GitHub: {self.resume.contact.github}")
        c.drawString(50, height - 140, f"LinkedIn: {self.resume.contact.linkedin}")

        # Add summary
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, height - 180, "Summary")
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 200, self.resume.summary)

        # Add objective
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, height - 240, "Objective")
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 260, self.resume.objective)

        # Add experience
        y_position = height - 300
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y_position, "Experience")
        y_position -= 20

        c.setFont("Helvetica", 12)
        for job in self.resume.experience:
            if y_position < 50:  # Check for page overflow
                c.showPage()
                y_position = height - 50

            c.drawString(50, y_position, f"{job.company} ({job.location})")
            y_position -= 20
            c.drawString(50, y_position, f"{job.title} ({job.dates})")
            y_position -= 20
            for responsibility in job.responsibilities:
                c.drawString(70, y_position, f"- {responsibility}")
                y_position -= 15

        # Add skills
        if y_position < 50:
            c.showPage()
            y_position = height - 50

        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y_position, "Skills")
        y_position -= 20

        c.setFont("Helvetica", 12)
        for skill_category, skills in self.resume.skills.model_dump().items():
            if skills:
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50

                c.drawString(50, y_position, skill_category.replace("_", " ").capitalize())
                y_position -= 20
                for skill in skills:
                    c.drawString(70, y_position, f"- {skill}")
                    y_position -= 15

        # Save the PDF
        c.save()